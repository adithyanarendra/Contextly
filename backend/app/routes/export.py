from fastapi import APIRouter, Depends, HTTPException
from pydantic import BaseModel
from sqlalchemy.orm import Session
from ..database import get_db
from ..models import QAHistory
from docx import Document as DocxDocument
from fastapi.responses import FileResponse
import tempfile
import os

router = APIRouter(prefix="/export", tags=["export"])


class ExportRequest(BaseModel):
    qa_ids: list[int]
    title: str = "exported_qa"


@router.post("/")
def export_qas(req: ExportRequest, db: Session = Depends(get_db)):
    if not req.qa_ids:
        raise HTTPException(status_code=400, detail="qa_ids is required")

    qas = db.query(QAHistory).filter(QAHistory.id.in_(req.qa_ids)).all()
    if not qas:
        raise HTTPException(status_code=404, detail="No QAs found for given ids")

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp.close()
    doc = DocxDocument()
    doc.add_heading(req.title, level=1)
    for q in qas:
        doc.add_heading(f"Q: {q.question}", level=2)
        doc.add_paragraph(f"A: {q.answer}")
        if q.source_chunk_ids:
            doc.add_paragraph(
                f"Sources chunk ids: {q.source_chunk_ids}", style="Intense Quote"
            )

    doc.save(tmp.name)
    return FileResponse(
        tmp.name,
        filename=f"{req.title}.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
